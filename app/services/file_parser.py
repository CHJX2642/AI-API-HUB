# -*- coding: utf-8 -*-
# AI API Hub — 文件文本提取模块
# 支持从 Word (.docx)、PDF (.pdf)、Excel (.xlsx) 文件中提取纯文本

import io                              # 内存文件流，用于将字节数据包装为文件对象
from docx import Document as DocxDocument   # Word 文档解析库
from PyPDF2 import PdfReader               # PDF 文本提取库
from openpyxl import load_workbook         # Excel 文件解析库


def extract_text_from_docx(file_bytes):
    """从 Word .docx 文件中提取文本（段落 + 表格）
    参数:
        file_bytes: 文件的二进制内容
    返回:
        提取的纯文本字符串，段落和表格内容用换行连接
    """
    doc = DocxDocument(io.BytesIO(file_bytes))   # 将字节流包装为文件对象并解析
    parts = []                                    # 存储所有文本片段

    # 提取段落文本
    for para in doc.paragraphs:                   # 遍历所有段落
        text = para.text.strip()                  # 去除首尾空白
        if text:                                  # 跳过空段落
            parts.append(text)

    # 提取表格文本（每行单元格用制表符分隔）
    for table in doc.tables:                      # 遍历所有表格
        for row in table.rows:                    # 遍历表格每行
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:                             # 跳过空行
                parts.append('\t'.join(cells))    # 单元格用制表符连接

    return '\n'.join(parts)                       # 所有片段用换行连接


def extract_text_from_pdf(file_bytes):
    """从 PDF 文件中提取文本（逐页提取）
    参数:
        file_bytes: 文件的二进制内容
    返回:
        提取的纯文本字符串，每页内容用双换行分隔
    """
    reader = PdfReader(io.BytesIO(file_bytes))    # 将字节流包装为文件对象并解析
    parts = []                                    # 存储每页文本

    for page in reader.pages:                     # 遍历每一页
        text = page.extract_text()                # 提取页面文本
        if text and text.strip():                 # 跳过空页
            parts.append(text.strip())

    return '\n\n'.join(parts)                     # 各页用双换行分隔


def extract_text_from_excel(file_bytes):
    """从 Excel .xlsx 文件中提取文本（制表符分隔）
    参数:
        file_bytes: 文件的二进制内容
    返回:
        提取的纯文本字符串，每个 sheet 有标题行，单元格用制表符分隔
    """
    # read_only=True 提升性能，data_only=True 只读取值不读公式
    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    parts = []                                    # 存储所有文本片段

    for sheet in wb.sheetnames:                   # 遍历所有工作表
        ws = wb[sheet]                            # 获取工作表对象
        parts.append(f'=== {sheet} ===')          # 添加工作表名称作为标题
        for row in ws.iter_rows(values_only=True): # 遍历每行（只取值）
            cells = [str(c) if c is not None else '' for c in row]  # 转为字符串
            if any(cells):                        # 跳过全空行
                parts.append('\t'.join(cells))    # 单元格用制表符连接

    wb.close()                                    # 关闭工作簿释放资源
    return '\n'.join(parts)                       # 所有片段用换行连接
